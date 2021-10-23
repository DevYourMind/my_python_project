from docx import Document
from docx.shared import Inches, Cm
from pytube import YouTube
from PIL import Image, ImageEnhance
from pytesseract import pytesseract
import glob
import cv2
import os


def get_title(videourl):
    '''
    Parameters
    ----------
    videourl : gets url

    Returns
    -------
    list
        title of video and stream of video for downoading
    '''
    yt = YouTube(videourl)
    return [yt.title, yt.streams.get_highest_resolution()]


def download_youtube(videourl, title):
    '''
    Function creates folder for video and downloads it.
    
    Parameters
    ----------
    videourl : url string
    title : title of video 
    
    ----------
    
    Returns
    -------
    None.
    
    '''
    yt = YouTube(videourl)
    path = "./" + title
    if not os.path.exists(path):
        os.makedirs(path)
    ys = yt.streams.get_highest_resolution()
    out_file = ys.download()
    os.rename(out_file, "./" + title + "/" + title + ".mp4")


def create_screenshots(title, interval, square):
    '''
    Creating, cropping and saving screenshots.
    
    Parameters
    ----------
    title : gets a video title for placing files in folder
    interval : interval of screenshots creating
    square : area of screenshot cropping

    Returns
    -------
    currentFrame : current number of frames

    '''
    cap = cv2.VideoCapture("./" + title + "/" + title + ".mp4")
    try:
        if not os.path.exists("./" + title + "/Screenshots"):
            os.makedirs(title + "/Screenshots")
    except OSError:
        print("Error: Creating directory of data")
    width = int(cap.get(3))
    height = int(cap.get(4))
    currentFrame = 0
    if len(square) > 1:
        ya = square[0]
        yb = square[1]
        xa = square[2]
        xb = square[3]
    else:
        width = int(cap.get(3))
        height = int(cap.get(4))
        ya = int(height * 0.54)
        yb = int(height)
        xa = 0
        xb = int(width)
    while True:
        for i in range(interval):
            ret, frame = cap.read()
        if ret == False:
            print("No Screenshots :(")
            break
        frame = frame[ya:yb, xa:xb]
        name = "./" + title + "/Screenshots/frame" + str(currentFrame) + ".jpg"
        cv2.imwrite(name, frame)
        currentFrame += 1
    cap.release()
    cv2.destroyAllWindows()
    return currentFrame


def delete_intro(num_of_frames, video_title, yellow=False):
    '''
    Deletes intro without notes in file.

    Parameters
    ----------
    num_of_frames : number of frames in file
    video_title : str with title
    yellow : special processing if video has yellow label. The default is False.

    Returns
    -------
    None.

    '''
    non_label_image = 0
    path_to_tesseract = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    pytesseract.tesseract_cmd = path_to_tesseract
    for i in range(num_of_frames):
        image_path = video_title + "/Screenshots/frame" + str(i) + ".jpg"
        img = Image.open(image_path)
        if yellow == False:
            img = cv2.imread(image_path)
            img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            text = pytesseract.image_to_string(img)
            non_label_image += 1
            os.remove(image_path)
            if ("SKY" in text) == False:
                break
        else:
            image_rgb = img.convert("RGB")
            rgb_pixel_value = image_rgb.getpixel((500, 10))
            if rgb_pixel_value == (0, 0, 0):
                non_label_image += 1
                os.remove(image_path)
    files = os.listdir("./" + video_title + "/Screenshots/")
    for file in files:
        os.rename(
            "./" + video_title + "/Screenshots/" + file,
            "./"
            + video_title
            + "/Screenshots/frame."
            + str(file.split("frame")[1].split(".jpg")[0])
            + ".jpg",
        )
    files = os.listdir("./" + video_title + "/Screenshots/")
    index = 0
    for file in files:
        os.rename(
            "./" + video_title + "/Screenshots/" + file,
            "./" + video_title + "/Screenshots/frame" + str(index) + ".jpg",
        )
        index += 1


def delete_duplicates(video_title, yellow=False):
    '''
    The function deletes duplicates of video using number
    in left corner.

    Parameters
    ----------
    video_title : TYPE
        DESCRIPTION.
    yellow : TYPE, optional
        DESCRIPTION. The default is False.

    Returns
    -------
    None.

    '''
    width, height = Image.open(video_title + "/Screenshots/frame0.jpg").size
    left = 0
    top = 5
    right = width / 10
    if yellow:
        bottom = height / 4.4
    else:
        bottom = height / 4.5
    path_to_tesseract = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    pytesseract.tesseract_cmd = path_to_tesseract
    num_of_files = len(os.listdir(video_title + "/Screenshots"))
    for i in range(1, num_of_files - 1):
        flag = False
        im1 = Image.open(video_title + "/Screenshots/frame" + str(i - 1) + ".jpg")
        im2 = Image.open(video_title + "/Screenshots/frame" + str(i) + ".jpg")
        im1 = im1.crop((left, top, right, bottom))
        enhancer = ImageEnhance.Contrast(im1)
        im1 = enhancer.enhance(2)
        im2 = im2.crop((left, top, right, bottom))
        enhancer = ImageEnhance.Contrast(im2)
        im2 = enhancer.enhance(2)
        im1.save(video_title + "/Screenshots/frame_temp" + str(i - 1) + ".jpg")
        im2.save(video_title + "/Screenshots/frame_temp" + str(i) + ".jpg")
        img1 = cv2.imread(video_title + "/Screenshots/frame_temp" + str(i - 1) + ".jpg")
        img2 = cv2.imread(video_title + "/Screenshots/frame_temp" + str(i) + ".jpg")
        img1 = cv2.cvtColor(img1, cv2.COLOR_BGR2GRAY)
        img2 = cv2.cvtColor(img2, cv2.COLOR_BGR2GRAY)
        text1 = pytesseract.image_to_string(img1)
        text2 = pytesseract.image_to_string(img2)
        try:
            t1 = int(text1[:-1])
            t2 = int(text2[:-1])
        except ValueError:
            print("Value Error")
        else:
            print(t1)
            print(t2)
            if int(text1[:-1]) == int(text2[:-1]):
                flag = True
        print("################")
        if flag:
            os.remove(video_title + "/Screenshots/frame" + str(i - 1) + ".jpg")
        os.remove(video_title + "/Screenshots/frame_temp" + str(i - 1) + ".jpg")
        for filename in glob.glob(video_title + "/Screenshots/frame_temp*"):
            os.remove(filename)


def create_word(title, currentFrame):
    '''
    Place screenshots in word file.

    Parameters
    ----------
    title : str with video title.
    currentFrame : integer of current number of screens.

    Returns
    -------
    None.

    '''
    document = Document()
    sections = document.sections
    files = os.listdir("./" + title + "/Screenshots/")
    for file in files:
        os.rename(
            "./" + title + "/Screenshots/" + file,
            "./"
            + title
            + "/Screenshots/frame."
            + str(file.split("frame")[1].split(".jpg")[0])
            + ".jpg",
        )
    files = os.listdir("./" + title + "/Screenshots/")
    sorted_files = sorted(files, key=lambda x: int(x.split(".")[1]))
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    files = glob.glob(title + "/Screenshots/frame*")
    for file in sorted_files:
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture("./" + title + "/Screenshots/" + file, width=Inches(8))
    document.save(title + "/" + title + ".docx")
    os.startfile(title)


def delete_everything(title):
    '''
    Delete all files in video folder for avoiding exceptions.

    Parameters
    ----------
    title : folder with name of video

    Returns
    -------
    None.

    '''
    top = "./" + title
    for root, dirs, files in os.walk(top, topdown=False):
        for name in files:
            os.remove(os.path.join(root, name))
        for name in dirs:
            os.rmdir(os.path.join(root, name))
